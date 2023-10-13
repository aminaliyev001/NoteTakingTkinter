from datetime import datetime
import tkinter as tk
from tkinter import ttk, messagebox
import json
from ttkbootstrap import Style
from PIL import Image, ImageTk
from tkinter import simpledialog
from docx import Document


WINDOW_WIDTH = 1000
WINDOW_HEIGHT = 500

def format_datetime(iso_str):
    try:
        dt_obj = datetime.fromisoformat(iso_str)
        return dt_obj.strftime('%Y-%m-%d %H:%M:%S')
    except ValueError:
        return iso_str
    
def save_subjects(subjects, filename="subjects.json"):
    with open(filename, 'w') as file:
        json.dump(subjects, file, indent=4) 

def load_subjects_from_file(filename="subjects.json"):
    try:
        with open(filename, 'r') as file:
            return json.load(file)
    except FileNotFoundError:
        return {}  
    except json.JSONDecodeError:
        print(f"Error: {filename} contains invalid JSON.")
        return {}

def load_notes_from_file(filename="notes.json"):
    notes = {}
    try:
        with open(filename, "r") as f:
            notes = json.load(f)
    except FileNotFoundError:
        messagebox.showwarning("Warning","File not found! Please Restart the application")
    return notes
def home_page(username):
    home_root = tk.Toplevel(root_login)
    home_root.title("MyNotes App Home")

    screen_width = home_root.winfo_screenwidth()
    screen_height = home_root.winfo_screenheight()
    x = (screen_width/2) - (WINDOW_WIDTH/2)
    y = (screen_height/2) - (WINDOW_HEIGHT/2)
    home_root.geometry('%dx%d+%d+%d' % (WINDOW_WIDTH,WINDOW_HEIGHT, x, y))

    welcome_label = ttk.Label(home_root, text=f"Welcome {username} to MyNotes App", font=("TkDefaultFont", 20, "bold"))
    welcome_label.pack(pady=20)
    
    bookshelf_frame = ttk.Frame(home_root)
    bookshelf_frame.pack(pady=20)
    
    subjects = load_subjects_from_file()
    def SubjectsExists(subject_name):
        for subject_id, subject in subjects.items():
            if(subject.lower() == subject_name.lower()):
                return True
        return False
    book_image = Image.open("book.png")
    book_photo = ImageTk.PhotoImage(book_image)
    notes_home = load_notes_from_file()
    for subject_id, subject in subjects.items():
        notes_count = len(notes_home.get(subject_id, {}).get("notes", {}))
        display_text = f"{subject}\n({notes_count} notes)"
        
        book_label = ttk.Button(
            bookshelf_frame, 
            text=display_text, 
            image=book_photo, 
            compound=tk.TOP, 
            command=lambda name=subject, sid=subject_id: main_page(home_root, sid)
        )
        book_label.pack(side=tk.LEFT, padx=10)

    def add_book():
        new_book = simpledialog.askstring("New Subject", "Enter the name of the new subject:")
        if new_book:
            if(SubjectsExists(new_book) == False):
                new_book_label = ttk.Button(bookshelf_frame, text=new_book, command=lambda: main_page(home_root))
                new_book_label.pack(side=tk.LEFT, padx=10)
                subjects.append(new_book)
                save_subjects(subjects)
            else:
                messagebox.showwarning("Warning",f"This subject name {new_book} already exists")
        elif new_book is not None:
            messagebox.showwarning("Warning","Subject could not be created! Fill the field")
    
    def delete_subject():
        subject_to_delete = simpledialog.askstring("Delete Subject", "Enter the name of the subject to delete:")
        subject_key = [key for key, value in subjects.items() if value.lower() == subject_to_delete.lower()]

        if subject_to_delete and subject_key:
            subject_key = subject_key[0]
            if SubjectsExists(subject_to_delete):
                response = messagebox.askyesno("Confirmation", f"Are you sure you want to delete {subject_to_delete}?")
                if response:
                    del subjects[subject_key]
                    save_subjects(subjects)
                    with open("notes.json", "r") as notes_file:
                        notes_data = json.load(notes_file)
                        if subject_key in notes_data:
                            del notes_data[subject_key]
                            with open("notes.json", "w") as notes_file_write:
                                json.dump(notes_data, notes_file_write)
                    
                    messagebox.showinfo("Success", f"{subject_to_delete} and its notes have been deleted!")
                    home_root.destroy()
                    home_page(username)
            else:
                messagebox.showwarning("Warning", f"The subject {subject_to_delete} does not exist.")
        elif subject_to_delete is not None:
            messagebox.showwarning("Warning", "Please enter a valid subject name.")


    def edit_subject():
        subject_to_edit = simpledialog.askstring("Edit Subject", "Enter the name of the subject to edit:")
        if subject_to_edit is None:
            return
        if not subject_to_edit.strip(): 
            messagebox.showwarning("Warning", "Subject could not be edited! Fill the field.")
            return
        if not SubjectsExists(subject_to_edit):  
            messagebox.showwarning("Warning", f"The subject {subject_to_edit} does not exist.")
            return
        new_subject_name = simpledialog.askstring("Edit Subject", f"Enter the new name for {subject_to_edit}:")
        if new_subject_name is None:
            return
        if not new_subject_name.strip():  
            messagebox.showwarning("Warning", "Subject could not be renamed! Fill the field.")
            return

        if SubjectsExists(new_subject_name):  
            messagebox.showwarning("Warning", f"The subject name {new_subject_name} already exists.")
            return

        for key, value in subjects.items():
            print(key, value)
            if value.lower() == subject_to_edit.strip().lower():
                subjects[key] = new_subject_name.strip() 
                print("yes")
                break
        save_subjects(subjects)
        messagebox.showinfo("Updated", f"Subject '{subject_to_edit}' has been updated to '{new_subject_name}'.")
        home_root.destroy()
        home_page(username)
    def destroyApp():
        home_root.destroy()
        root_login.destroy()
        
    button_frame = ttk.Frame(home_root)
    button_frame.pack(side=tk.BOTTOM, pady=20)  

    add_button = ttk.Button(button_frame, text="Add New Subject", command=add_book, style="success.TButton")
    add_button.pack(side=tk.LEFT, padx=10)  

    edit_button = ttk.Button(button_frame, text="Edit a Subject", command=edit_subject, style="primary.TButton")
    edit_button.pack(side=tk.LEFT, padx=10)

    delete_button = ttk.Button(button_frame, text="Delete Subject", command=delete_subject, style="danger.TButton")
    delete_button.pack(side=tk.LEFT, padx=10)

    logout_button = ttk.Button(button_frame, text="Logout", command=destroyApp, style="dark.TButton")
    logout_button.pack(side=tk.LEFT, padx=10)

    home_root.mainloop()

def main_page(root_login,id):
    root = tk.Toplevel(root_login)
    root.title("MyNotes App")
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x = (screen_width/2) - (WINDOW_WIDTH/2)
    y = (screen_height/2) - (WINDOW_HEIGHT/2)
    root.geometry('%dx%d+%d+%d' % (WINDOW_WIDTH,WINDOW_HEIGHT, x, y))
    style.configure("TNotebook.Tab", font=("TkDefaultFont", 14, "bold"))
    notebook = ttk.Notebook(root, style="TNotebook")
    notes = load_notes_from_file()
    fileName = "notes.json"

    notebook = ttk.Notebook(root)
    notebook.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

    def export_as_docx(title, content):
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(content)
        
        file_path = f"{title}.docx"
        doc.save(file_path)
        
        messagebox.showinfo("Success", f"'{title}' has been saved as a .docx file!")
    def back():
        root.destroy()
    def update_note():
        current_tab = notebook.index(notebook.select())
        note_title = notebook.tab(current_tab, "text")
        note_frame = notebook.nametowidget(notebook.select())
        for child in note_frame.winfo_children():
            if isinstance(child, tk.Text):
                note_content_widget = child
                break
        note_content = note_content_widget.get("1.0", tk.END).strip()
        if id not in notes:
            notes[id] = {"notes": {}}
        
        if note_title in notes[id]["notes"]:
            notes[id]["notes"][note_title]["content"] = note_content
            notes[id]["notes"][note_title]["char_count"] = len(note_content)
            notes[id]["notes"][note_title]["last_modified"] = datetime.utcnow().isoformat()
        else:
            notes[id]["notes"][note_title] = {
                "content": note_content,
                "created": datetime.now().isoformat(),
                "last_modified": datetime.now().isoformat(),
                "char_count": len(note_content)
            }

        with open(fileName, "w") as f:
            json.dump(notes, f)
        messagebox.showinfo("Updated", f"Note '{note_title}' has been updated.")


    def add_note():
        note_frame = ttk.Frame(notebook, padding=10)
        notebook.add(note_frame, text="New Note")
        title_label = ttk.Label(note_frame, text="Title:")
        title_label.grid(row=0, column=0, padx=10, pady=10, sticky="W")
        title_entry = ttk.Entry(note_frame, width=40)
        title_entry.grid(row=0, column=1, padx=10, pady=10)
        content_label = ttk.Label(note_frame, text="Content:")
        content_label.grid(row=1, column=0, padx=10, pady=10, sticky="W")
        content_entry = tk.Text(note_frame, width=40, height=10)
        content_entry.grid(row=1, column=1, padx=10, pady=10)
        def save_note():
            title = title_entry.get()
            content = content_entry.get("1.0", tk.END).strip()
            if not title or not content:
                messagebox.showwarning("Warning", "Both title and content should be filled!")
                return
            if id not in notes:
                notes[id] = {"notes": {}}

            created_time = datetime.now().isoformat()
            notes[id]["notes"][title] = {
                "content": content,
                "created": created_time,
                "last_modified": created_time,
                "char_count": len(content)
            }
            with open(fileName, "w") as f:
                json.dump(notes, f)

            root.destroy()
            main_page(root_login,id)

        save_button = ttk.Button(note_frame, text="Save", command=save_note, style="secondary.TButton")
        save_button.grid(row=2, column=1, padx=10, pady=10)

    def load_notes():
        try:
            with open(fileName, "r") as f:
                notes = json.load(f)

            if id in notes:
                for note_title, content_data in notes[id]["notes"].items():
                    note_frame = ttk.Frame(notebook)
                    note_content = tk.Text(note_frame, width=40, height=10)
                    note_content.grid(row=0, column=0, sticky="nsew")
                    note_frame.grid_columnconfigure(0, weight=1)
                    note_frame.grid_rowconfigure(0, weight=1)
                    if "content" in content_data:
                        note_content.insert(tk.END, content_data["content"])
                        char_count = content_data.get("char_count", 0)
                        created_time = format_datetime(content_data.get("created", "N/A"))
                        last_modified = format_datetime(content_data.get("last_modified", "N/A"))
                        info_label = ttk.Label(note_frame, text=f"Character count: {char_count} | Created: {created_time} | Last modified: {last_modified}")
                        info_label.grid(row=1, column=0, sticky="w")
                    else:
                        note_content.insert(tk.END, "No content available.")
                    notebook.add(note_frame, text=note_title)
            else:
                note_content = tk.Text(notebook, width=40, height=10)
                note_content.insert(tk.END, "No notes available for this subject.")
                notebook.add(note_content, text="Empty")

        except FileNotFoundError:
            messagebox.showwarning("Warning","File not found! Please restart the application")
            root.destroy()

            
    load_notes()
    def delete_note():
        current_tab = notebook.index(notebook.select())
        note_title = notebook.tab(current_tab, "text")
        confirm = messagebox.askyesno("Delete Note", f"Are you sure you want to delete {note_title}?")
        
        if confirm:
            notebook.forget(current_tab)
            if note_title in notes[id]["notes"]:
                notes[id]["notes"].pop(note_title)
                with open(fileName, "w") as f:
                    json.dump(notes, f)
    def change_title():
        current_tab = notebook.index(notebook.select())
        current_title = notebook.tab(current_tab, "text")
        WINDOW_WIDTH_CT = 300
        WINDOW_HEIGHT_CT = 150
        new_window = tk.Toplevel(root)
        new_window.title("Change Title")
        screen_width = new_window.winfo_screenwidth()
        screen_height = new_window.winfo_screenheight()
        x = (screen_width/2) - (WINDOW_WIDTH_CT/2)
        y = (screen_height/2) - (WINDOW_HEIGHT_CT/2)
        new_window.geometry('%dx%d+%d+%d' % (WINDOW_WIDTH_CT,WINDOW_HEIGHT_CT, x, y))

        new_title_label = ttk.Label(new_window, text="New Title:")
        new_title_label.pack(pady=10, padx=10, anchor="w")

        new_title_entry = ttk.Entry(new_window, width=30)
        new_title_entry.pack(padx=10, pady=5)
        new_title_entry.insert(0, current_title) 

        def save_new_title():
            new_title = new_title_entry.get().strip()
            notebook.tab(current_tab, text=new_title)
            
            if current_title in notes[id]["notes"]:
                notes[id]["notes"][new_title] = notes[id]["notes"].pop(current_title)
                with open(fileName, "w") as f:
                    json.dump(notes, f)
                new_window.destroy()

        save_button = ttk.Button(new_window, text="Save", command=save_new_title)
        save_button.pack(pady=10)

    new_button = ttk.Button(root, text="Update", 
                            command=update_note, style="success.TButton")
    new_button.pack(side=tk.LEFT, padx=10, pady=10)

    new_button = ttk.Button(root, text="New Note", 
                            command=add_note, style="info.TButton")
    new_button.pack(side=tk.LEFT, padx=10, pady=10)

    change_title_button = ttk.Button(root, text="Change Title", command=change_title, style="warning.TButton")
    change_title_button.pack(side=tk.LEFT, padx=10, pady=10)

    delete_button = ttk.Button(root, text="Delete", 
                            command=delete_note, style="primary.TButton")
    delete_button.pack(side=tk.LEFT, padx=10, pady=10)

    def export_current_note():
        current_tab = notebook.index(notebook.select())
        note_title = notebook.tab(current_tab, "text")
        frame = notebook.nametowidget(notebook.select())
        text_widget = None
        for child in frame.winfo_children():
            if isinstance(child, tk.Text):
                text_widget = child
                break
        if text_widget:
            note_content = text_widget.get("1.0", tk.END).strip()
            export_as_docx(note_title, note_content)
        else:
            messagebox.showerror("Error", "Could not fetch note content!")

    exit_button = ttk.Button(root,text="Back",command=back,style="dark.TButton")
    exit_button.pack(side = tk.RIGHT,padx=10,pady=10)
    
    export_button = ttk.Button(root, text="Export as docx", command=export_current_note)
    export_button.pack(side=tk.RIGHT, pady=10,padx=10) 

    root.mainloop()
    
def load_users():
    try:
        with open("users.json", "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return {}
def login():
    users = load_users()
    username = username_entry.get()
    password = password_entry.get()

    if username in users and users[username] == password:
        login_frame.pack_forget()
        root_login.withdraw()
        home_page(username)
    else:
        messagebox.showwarning("Login Failed", "Incorrect username or password")
        
root_login = tk.Tk()
root_login.title("Login - MyNotes App")
style = Style(theme='journal')

screen_width = root_login.winfo_screenwidth()
screen_height = root_login.winfo_screenheight()
x = (screen_width/2) - (WINDOW_WIDTH/2)
y = (screen_height/2) - (WINDOW_HEIGHT/2)
root_login.geometry('%dx%d+%d+%d' % (WINDOW_WIDTH,WINDOW_HEIGHT, x, y))

image = Image.open("bg.jpg")
image = image.resize((1000, 500), Image.ANTIALIAS)
photo = ImageTk.PhotoImage(image)
background_label = tk.Label(root_login, image=photo)
background_label.place(x=0, y=0, relwidth=1, relheight=1)
background_label.lower()

login_frame = ttk.Frame(root_login)
login_frame.pack(pady=150)

username_label = ttk.Label(login_frame, text="Username")
username_label.grid(row=0, column=0, padx=10, pady=10)

username_entry = ttk.Entry(login_frame, width=30)
username_entry.grid(row=0, column=1, padx=10, pady=10)

password_label = ttk.Label(login_frame, text="Password")
password_label.grid(row=1, column=0, padx=10, pady=10)

password_entry = ttk.Entry(login_frame, show="*", width=30)
password_entry.grid(row=1, column=1, padx=10, pady=10)

login_button = ttk.Button(login_frame, text="Login", command=login)
login_button.grid(row=2, column=0, columnspan=2, pady=20)
root_login.mainloop()